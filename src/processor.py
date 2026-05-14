from pathlib import Path
from datetime import datetime
import argparse
from collections import defaultdict

from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter


def _extract_docx_text(file_path: Path) -> str:
	"""Extract DOCX text while preserving paragraph/table order."""
	try:
		import docx  # python-docx
		from docx.table import Table
		from docx.text.paragraph import Paragraph
	except ImportError as exc:
		raise ImportError(
			"DOCX support requires 'python-docx'. "
			"Install with: pip install python-docx"
		) from exc

	doc = docx.Document(str(file_path))
	parts = []

	for child in doc.element.body.iterchildren():
		if child.tag.endswith("}p"):
			paragraph = Paragraph(child, doc)
			text = (paragraph.text or "").strip()
			if text:
				parts.append(text)
		elif child.tag.endswith("}tbl"):
			table = Table(child, doc)
			for row in table.rows:
				cells = [" ".join((cell.text or "").split()) for cell in row.cells]
				cells = [cell for cell in cells if cell]
				if cells:
					parts.append(" | ".join(cells))

	return "\n".join(parts).strip()


def _load_docx(file_path: Path):
	"""Load a DOCX file into a list of LangChain Document objects."""
	content = _extract_docx_text(file_path)
	if not content:
		raise ValueError(f"No extractable text found in DOCX file: {file_path}")
	return [Document(page_content=content, metadata={"source": str(file_path), "page": None})]


def load_document(file_paths):
	"""Load one or many files and return a flat list of Document objects."""
	if isinstance(file_paths, (str, Path)):
		paths = [file_paths]
	else:
		paths = list(file_paths)

	all_documents = []

	for raw_path in paths:
		path = Path(raw_path)
		suffix = path.suffix.lower()

		if suffix == ".pdf":
			documents = PDFPlumberLoader(str(path)).load()
		elif suffix == ".docx":
			documents = _load_docx(path)
		else:
			raise ValueError(f"Unsupported file format: {suffix}")

		date_uploaded = datetime.now().isoformat()
		for document in documents:
			metadata = dict(document.metadata or {})
			original_page = metadata.get("page")
			
			# Standardize page format: always int or None (never string)
			page_value = None
			if original_page is not None:
				try:
					page_value = int(original_page)
				except (TypeError, ValueError):
					page_value = None
			
			metadata.update(
				{
					"source": path.name,
					"source_path": str(path.resolve()),
					"file_type": suffix.lstrip("."),   
					"date_uploaded": date_uploaded,
					"page": page_value,  # Always int or None
				}
			)
			document.metadata = metadata
			all_documents.append(document)

	return all_documents


def split_text(documents, chunk_size=1000, chunk_overlap=100):
	"""Split documents while preserving source metadata in every chunk."""
	if chunk_size <= 0:
		raise ValueError("chunk_size must be greater than 0")
	if chunk_overlap < 0:
		raise ValueError("chunk_overlap must be greater than or equal to 0")
	if chunk_overlap >= chunk_size:
		raise ValueError("chunk_overlap must be smaller than chunk_size")

	text_splitter = RecursiveCharacterTextSplitter(
		chunk_size=chunk_size,
		chunk_overlap=chunk_overlap,
	)

	chunks = []
	for document in documents:
		base_metadata = dict(document.metadata or {})
		base_metadata.update(
			{
				"chunk_size": chunk_size,
				"chunk_overlap": chunk_overlap,
			}
		)
		cursor = 0
		for chunk_index, chunk_text in enumerate(text_splitter.split_text(document.page_content), 1):
			start = document.page_content.find(chunk_text, cursor)
			if start < 0:
				start = document.page_content.find(chunk_text)
			if start < 0:
				start = None
				end = None
			else:
				end = start + len(chunk_text)
				cursor = max(start + 1, end - chunk_overlap)

			chunk_metadata = dict(base_metadata)
			chunk_metadata.update(
				{
					"chunk_index": chunk_index,
					"chunk_id": (
						f"{base_metadata.get('source', 'unknown')}:"
						f"{base_metadata.get('page', 'na')}:"
						f"{chunk_index}"
					),
					"char_start": start,
					"char_end": end,
				}
			)
			chunks.append(Document(page_content=chunk_text, metadata=chunk_metadata))

	return chunks


def get_embedding_model():
	"""Create and return a configured HuggingFace embedding model."""
	return HuggingFaceEmbeddings(
		model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2",
		model_kwargs={"device": "cpu"},
		encode_kwargs={"normalize_embeddings": True},
	)


def process_multiple_documents(file_paths, chunk_size=1000, chunk_overlap=100):
	"""Process multiple files and return final chunks ready for vector database ingestion."""
	documents = load_document(file_paths)
	return split_text(documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def process_pipeline(file_path, chunk_size=1000, chunk_overlap=100):
	"""Run the full processing pipeline and return chunks with embedding model."""
	docs = load_document(file_path)
	chunks = split_text(docs, chunk_size, chunk_overlap)
	return {
		"chunks": [chunk.page_content for chunk in chunks],
		"embedding_model": get_embedding_model(),
	}

#Hàm để test file docx và pdf, nếu không có file nào được cung cấp thì sẽ tạo 2 file docx mẫu để test
#chỉ để test, không phải là phần chính của module, nên sẽ xóa sau khi test xong
if __name__ == "__main__":
	parser = argparse.ArgumentParser(description="Test document processing pipeline")
	parser.add_argument(
		"files",
		nargs="*",
		help="Paths to input files (.pdf/.docx). If omitted, creates sample DOCX files.",
	)
	parser.add_argument("--chunk-size", type=int, default=500)
	parser.add_argument("--chunk-overlap", type=int, default=50)
	args = parser.parse_args()

	created_test_files = []
	input_files = []

	if args.files:
		input_files = [Path(p) for p in args.files]
	else:
		# Create 2 sample DOCX files for multi-file testing.
		try:
			import docx  # python-docx
		except ImportError as exc:
			raise ImportError(
				"To run the built-in DOCX test, install: pip install python-docx"
			) from exc

		base_dir = Path(__file__).parent
		for idx in range(1, 3):
			p = base_dir / f"test_{idx}.docx"
			d = docx.Document()
			d.add_paragraph(f"Tai lieu DOCX test so {idx}. " * 120)
			d.save(str(p))
			created_test_files.append(p)
			input_files.append(p)

	success = False
	try:
		chunks = process_multiple_documents(
			[str(p) for p in input_files],
			chunk_size=args.chunk_size,
			chunk_overlap=args.chunk_overlap,
		)
		print(f"Tong so chunks duoc tao: {len(chunks)}")
		if not chunks:
			print("Khong co chunk nao duoc tao.")
		else:
			counts_by_source = defaultdict(int)
			first_chunk_by_source = {}
			for chunk in chunks:
				source = (chunk.metadata or {}).get("source", "<unknown>")
				counts_by_source[source] += 1
				if source not in first_chunk_by_source:
					first_chunk_by_source[source] = chunk

			print("So chunks theo tung file:")
			for source in sorted(counts_by_source.keys()):
				print(f"- {source}: {counts_by_source[source]}")

			print("\nChunk dau tien cua moi file:")
			for source in sorted(first_chunk_by_source.keys()):
				chunk = first_chunk_by_source[source]
				print(f"\n=== {source} ({chunk.metadata.get('file_type')}) ===")
				print("Metadata:")
				print(chunk.metadata)
				print("Noi dung (doan dau):")
				print(chunk.page_content)
		success = True
	finally:
		if success:
			for p in created_test_files:
				if p.exists():
					p.unlink()
					print(f"Da xoa file test: {p.name}")
